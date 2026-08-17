import csv
import re
from io import BytesIO, StringIO

from docx import Document
from pypdf import PdfReader


class KnowledgeProcessingError(ValueError):
    pass


def extract_text(content: bytes, file_type: str) -> str:
    normalized_type = file_type.lower().lstrip(".")
    try:
        if normalized_type == "pdf":
            reader = PdfReader(BytesIO(content))
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        elif normalized_type == "docx":
            document = Document(BytesIO(content))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            table_rows = [
                " | ".join(cell.text.strip() for cell in row.cells)
                for table in document.tables
                for row in table.rows
            ]
            if table_rows:
                text = f"{text}\n" + "\n".join(table_rows)
        elif normalized_type == "csv":
            decoded = content.decode("utf-8-sig")
            text = "\n".join(" | ".join(row) for row in csv.reader(StringIO(decoded)))
        elif normalized_type in {"txt", "md"}:
            text = content.decode("utf-8-sig")
        else:
            raise KnowledgeProcessingError(f"Unsupported file type: {file_type}")
    except (UnicodeDecodeError, OSError, ValueError) as exc:
        raise KnowledgeProcessingError("The document could not be parsed") from exc

    cleaned = re.sub(r"[ \t]+", " ", text)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    if not cleaned:
        raise KnowledgeProcessingError("The document contains no extractable text")
    return cleaned


def chunk_text(text: str, *, size: int, overlap: int) -> list[str]:
    if size < 200 or overlap < 0 or overlap >= size:
        raise ValueError("Invalid knowledge chunk configuration")
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        remaining = paragraph
        while remaining:
            capacity = size - len(current) - (2 if current else 0)
            if capacity <= 0:
                chunks.append(current.strip())
                current = current[-overlap:].lstrip() if overlap else ""
                continue
            if len(remaining) <= capacity:
                current = f"{current}\n\n{remaining}".strip()
                remaining = ""
                continue
            split_at = remaining.rfind(" ", 0, capacity)
            if split_at < max(80, capacity // 2):
                split_at = capacity
            current = f"{current}\n\n{remaining[:split_at]}".strip()
            chunks.append(current)
            prefix = current[-overlap:].lstrip() if overlap else ""
            current = prefix
            remaining = remaining[split_at:].lstrip()
    if current.strip():
        chunks.append(current.strip())
    return chunks
